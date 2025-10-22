import sys
sys.dont_write_bytecode = True  # 🚫 Отключаем создание __pycache__ и .pyc файлов

import os
import re
from docx import Document
import pandas as pd
from tkinter import Tk, filedialog
from openpyxl import load_workbook
from openpyxl.styles import Border, Side, Alignment, PatternFill

# ======================================================
# 🌍 ОБЩИЕ ФУНКЦИИ
# ======================================================

TRANSLIT_TABLE = str.maketrans({
    "А": "A", "Б": "B", "В": "V", "Г": "G", "Д": "D",
    "Е": "E", "Ё": "YO", "Ж": "ZH", "З": "Z", "И": "I",
    "Й": "Y", "К": "K", "Л": "L", "М": "M", "Н": "N",
    "О": "O", "П": "P", "Р": "R", "С": "S", "Т": "T",
    "У": "U", "Ф": "F", "Х": "KH", "Ц": "TS", "Ч": "CH",
    "Ш": "SH", "Щ": "SHCH", "Ъ": "", "Ы": "Y", "Ь": "",
    "Э": "E", "Ю": "YU", "Я": "YA"
})

def transliterate(text: str) -> str:
    if not text:
        return ""
    return " ".join([p.upper().translate(TRANSLIT_TABLE).capitalize() for p in text.split()])

def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\xa0", " ")).strip()

def format_excel(filepath):
    wb = load_workbook(filepath)
    ws = wb.active

    border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    yellow = PatternFill(start_color="FFF9C4", end_color="FFF9C4", fill_type="solid")

    headers = [c.value for c in ws[1]]
    surname_col = headers.index("Фамилия (англ)") + 1 if "Фамилия (англ)" in headers else None
    name_col = headers.index("Имя (англ)") + 1 if "Имя (англ)" in headers else None

    for row in ws.iter_rows():
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            if cell.column in [surname_col, name_col]:
                cell.fill = yellow

    for col in ws.columns:
        max_len = max(len(str(c.value)) if c.value else 0 for c in col)
        ws.column_dimensions[col[0].column_letter].width = max_len + 3

    wb.save(filepath)

def extract_text_from_docx(doc):
    text = []
    for p in doc.paragraphs:
        if p.text.strip():
            text.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text.append(cell.text)
    return clean_text(" ".join(text))

# ======================================================
# 🇲🇲 ПАРСЕР МЬЯНМА
# ======================================================

MONTHS = {
    "ЯНВАРЯ": "01", "ФЕВРАЛЯ": "02", "МАРТА": "03", "АПРЕЛЯ": "04",
    "МАЯ": "05", "ИЮНЯ": "06", "ИЮЛЯ": "07", "АВГУСТА": "08",
    "СЕНТЯБРЯ": "09", "ОКТЯБРЯ": "10", "НОЯБРЯ": "11", "ДЕКАБРЯ": "12"
}

def convert_date(date_str: str) -> str:
    if not date_str:
        return ""
    match = re.search(r"(\d{1,2})\s+([А-Яа-я]+)\s+(\d{4})", date_str)
    if match:
        day, month_word, year = match.groups()
        month = MONTHS.get(month_word.upper(), "??")
        return f"{int(day):02d}.{month}.{year}"
    return ""

def extract_myanmar_data(doc_path, log):
    try:
        doc = Document(doc_path)
        text = extract_text_from_docx(doc)
        data = {"Файл": os.path.basename(doc_path), "Страна": "Мьянма"}

        name_match = re.search(r"Имя\s+([А-ЯA-ZЁІЇЄа-яёіїє\s]+?)\s+Гражданство", text, re.IGNORECASE)
        name_ru = clean_text(name_match.group(1)) if name_match else ""
        name_en = transliterate(name_ru)

        birth_match = re.search(r"Дата рождения\s+(\d{1,2}\s+[А-Яа-я]+\s+\d{4})", text)
        birth_date = convert_date(birth_match.group(1)) if birth_match else ""

        gender_match = re.search(r"Пол\s*([МЖ])", text)
        gender = gender_match.group(1) if gender_match else ""
        birthplace_match = re.search(r"Место рождения\s+([A-Za-zА-Яа-яЁё\s,]+?)\s+Дата выдачи", text)
        birthplace = re.sub(r"^[МЖ]\s+", "", birthplace_match.group(1)).strip() if birthplace_match else ""

        passport_match = re.search(r"\b([A-ZА-Я]{1,2}\d{6,})\b", text)
        passport_number = passport_match.group(1).upper() if passport_match else ""

        issued_match = re.search(r"Орган выдачи\s+(.+?)\s+Действителен до", text)
        issued_raw = clean_text(issued_match.group(1)) if issued_match else ""
        issued_raw = issued_raw.replace("Действителен до", "").strip()

        date_match = re.search(r"(\d{1,2}\s+[А-Яа-я]+\s+\d{4})", issued_raw)
        issue_date = convert_date(date_match.group(1)) if date_match else ""
        issued_by = re.sub(r"\d{1,2}\s+[А-Яа-я]+\s+\d{4}", "", issued_raw).strip(" ,")

        expiry_match = re.search(r"Действителен до\s+(\d{1,2}\s+[А-Яа-я]+\s+\d{4})", text)
        expiry_date = convert_date(expiry_match.group(1)) if expiry_match else ""

        data.update({
            "Фамилия (рус)": "",
            "Фамилия (англ)": "",
            "Имя (рус)": name_ru,
            "Имя (англ)": name_en,
            "Дата рождения": birth_date,
            "Пол": gender,
            "Номер паспорта": passport_number,
            "Дата выдачи": issue_date,
            "Действителен до": expiry_date,
            "Место рождения": birthplace,
            "Кем выдан": issued_by
        })
        return data
    except Exception as e:
        log.append(f"Ошибка при обработке {os.path.basename(doc_path)} (Мьянма): {e}")
        return None

# ======================================================
# 🇹🇲 ПАРСЕР ТУРКМЕНИСТАН
# ======================================================

def extract_turkmenistan_data(doc_path, log):
    try:
        from parser_v5 import surname_dict, name_dict
    except:
        surname_dict = {}
        name_dict = {}

    try:
        doc = Document(doc_path)
        lines = []
        for p in doc.paragraphs:
            t = clean_text(p.text)
            if t:
                lines.append(t)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = clean_text(cell.text)
                    if t:
                        lines.append(t)

        def get_value(lines, key):
            for i, line in enumerate(lines):
                if key.lower() in line.lower():
                    for j in range(i + 1, len(lines)):
                        val = clean_text(lines[j])
                        if val:
                            return val
            return ""

        data = {"Файл": os.path.basename(doc_path), "Страна": "Туркменистан"}

        passport_number = ""
        for i, line in enumerate(lines):
            if "номер паспорта" in line.lower():
                window = lines[max(0, i - 3): min(len(lines), i + 4)]
                for segment in window:
                    match = re.search(r"[A-ZА-Я]\d{6,}", segment)
                    if match:
                        passport_number = match.group(0)
                        break
                break
        if not passport_number:
            for line in lines:
                match = re.search(r"[A-ZА-Я]\d{6,}", line)
                if match:
                    passport_number = match.group(0)
                    break

        surname_ru = get_value(lines, "Фамилия")
        name_ru = get_value(lines, "Имя")

        surname_en = surname_dict.get(surname_ru, transliterate(surname_ru))
        name_en = name_dict.get(name_ru, transliterate(name_ru))

        combined = get_value(lines, "Дата рождения")
        date_match = re.search(r"\d{2}\.\d{2}\.\d{4}", combined)
        gender = ""
        if date_match:
            birth_date = date_match.group(0)
            rest = combined.replace(birth_date, "").strip()
            if "муж" in rest.lower():
                gender = "М"
            elif "жен" in rest.lower():
                gender = "Ж"
        else:
            birth_date = combined

        issued_by = get_value(lines, "Орган,  выдавший документ") or get_value(lines, "Орган, выдавший документ")
        issued_by = issued_by.strip()

        if re.search(r"\bгмст\b", issued_by.lower()) or "миграц" in issued_by.lower():
            issued_by = "Государственная миграционная служба Туркменистана"

        data.update({
            "Фамилия (рус)": surname_ru,
            "Фамилия (англ)": surname_en,
            "Имя (рус)": name_ru,
            "Имя (англ)": name_en,
            "Дата рождения": birth_date,
            "Пол": gender,
            "Номер паспорта": passport_number,
            "Дата выдачи": get_value(lines, "Дата выдачи"),
            "Действителен до": get_value(lines, "Действителен до"),
            "Место рождения": get_value(lines, "Место рождения"),
            "Кем выдан": issued_by
        })
        return data
    except Exception as e:
        log.append(f"Ошибка при обработке {os.path.basename(doc_path)} (Туркменистан): {e}")
        return None

# ======================================================
# ✨ НОРМАЛИЗАЦИЯ РЕГИСТРА
# ======================================================

def normalize_data(df):
    def capitalize_text(x, column_name):
        if isinstance(x, str):
            if column_name == "Номер паспорта":
                return x.strip().upper()
            return " ".join(word.capitalize() for word in x.strip().split())
        return x

    for col in df.columns:
        df[col] = df[col].apply(lambda v: capitalize_text(v, col))
    return df

# ======================================================
# 🚀 ГЛАВНЫЙ ЗАПУСК
# ======================================================

def main():
    Tk().withdraw()
    folder = filedialog.askdirectory(title="Выберите папку с паспортами")
    if not folder:
        print("🚫 Папка не выбрана.")
        return

    results, log = [], []

    for file in os.listdir(folder):
        if not file.lower().endswith(".docx"):
            continue
        file_path = os.path.join(folder, file)
        try:
            doc = Document(file_path)
            text = extract_text_from_docx(doc)
        except:
            continue

        if "мьянм" in text.lower() or "myanmar" in text.lower():
            data = extract_myanmar_data(file_path, log)
        else:
            data = extract_turkmenistan_data(file_path, log)

        if data:
            results.append(data)

    if not results:
        print("❌ Не найдено документов.")
        return

    df = pd.DataFrame(results)
    cols = ["Фамилия (рус)", "Фамилия (англ)", "Имя (рус)", "Имя (англ)",
            "Дата рождения", "Номер паспорта", "Дата выдачи", "Действителен до",
            "Место рождения", "Кем выдан", "Страна", "Пол"]
    df = df.reindex(columns=cols)

    df = normalize_data(df)

    out = os.path.join(os.path.expanduser("~"), "Desktop", "ИТОГ.xlsx")
    df.to_excel(out, index=False)
    format_excel(out)

    print(f"✅ Таблица успешно создана: {out}")
    if log:
        print("⚠️ Логи ошибок:")
        for e in log:
            print(e)

if __name__ == "__main__":
    main()
